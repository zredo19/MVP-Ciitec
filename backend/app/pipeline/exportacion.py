"""
Filtro de salida — EXPORTACIÓN (RF-008): briefing -> PDF / Word / texto plano
con la plantilla institucional de 6 páginas.

- PDF (maestro): Jinja2 -> HTML+CSS -> WeasyPrint. Gráficos (donut/barras) se
  renderizan server-side con matplotlib a PNG embebidos (data URI).
- Word: documento .docx generado con python-docx (mismo contenido por secciones).
- Texto: volcado estructurado.
"""
from __future__ import annotations

import base64
import io
import json
import os
import re
from datetime import datetime
from functools import lru_cache
from typing import Any

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt  # noqa: E402
from jinja2 import Environment, FileSystemLoader, select_autoescape  # noqa: E402

from ..config import settings  # noqa: E402

_TPL_DIR = os.path.join(os.path.dirname(__file__), "..", "templates")
_env = Environment(loader=FileSystemLoader(_TPL_DIR), autoescape=select_autoescape(["html", "xml"]))

# Paleta institucional (consistente con la plantilla).
_AZUL = "#2c3e50"
_AZUL2 = "#6f86b3"


def _num(v: Any) -> float | None:
    """Convierte '2.305', '5.911', 50 -> float; '-.-'/None -> None."""
    if v is None:
        return None
    s = str(v).strip().replace(".", "").replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def _png_data_uri(fig) -> str:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=110, bbox_inches="tight", transparent=True)
    plt.close(fig)
    return "data:image/png;base64," + base64.b64encode(buf.getvalue()).decode("ascii")


def _donut_personal(personal: dict) -> str | None:
    pf = (personal or {}).get("parte_fuerza_institucional") or {}
    datos = {k: _num(pf.get(k)) for k in ("OF", "SOF", "ECP", "SLTP")}
    datos = {k: v for k, v in datos.items() if v}
    if not datos:
        return None
    fig, ax = plt.subplots(figsize=(2.6, 2.6))
    ax.pie(list(datos.values()), labels=list(datos.keys()), wedgeprops={"width": 0.42},
           colors=[_AZUL, _AZUL2, "#9fb0cc", "#c9d3e3"], textprops={"fontsize": 8})
    ax.set(aspect="equal")
    return _png_data_uri(fig)


def _bar_macrozonas(personal: dict) -> str | None:
    mz = (personal or {}).get("macrozonas") or []
    pares = [(m.get("zona", ""), _num(m.get("fuerza"))) for m in mz]
    pares = [(z, v) for z, v in pares if v]
    if not pares:
        return None
    fig, ax = plt.subplots(figsize=(3.4, 2.4))
    ax.barh([z for z, _ in pares], [v for _, v in pares], color=_AZUL)
    ax.invert_yaxis()
    ax.tick_params(labelsize=7)
    return _png_data_uri(fig)


def _bar_traslados(operaciones: dict) -> str | None:
    tr = (operaciones or {}).get("traslados_uac") or []
    pares = [(t.get("uac", ""), _num(t.get("unidades"))) for t in tr]
    pares = [(u, v) for u, v in pares if v]
    if not pares:
        return None
    fig, ax = plt.subplots(figsize=(4.2, 2.2))
    ax.bar([u for u, _ in pares], [v for _, v in pares], color=_AZUL2)
    ax.tick_params(labelsize=6, axis="x", rotation=45)
    return _png_data_uri(fig)


# ---------------------------------------------------------------------------
#  Modo PIXEL-PERFECT (overlay sobre la plantilla institucional en PNG)
# ---------------------------------------------------------------------------
@lru_cache(maxsize=1)
def _overlay_map() -> dict:
    with open(os.path.join(_TPL_DIR, "overlay.json"), encoding="utf-8") as fh:
        return json.load(fh)


@lru_cache(maxsize=8)
def _png_uri(num: int) -> str | None:
    ruta = os.path.join(_TPL_DIR, f"{num}.PNG")
    if not os.path.exists(ruta):
        return None
    with open(ruta, "rb") as fh:
        return "data:image/png;base64," + base64.b64encode(fh.read()).decode("ascii")


def _resolver(contenido: dict, ruta: str):
    """Resuelve rutas tipo 'a.b[0].c' dentro del contenido."""
    cur: Any = contenido
    for parte in ruta.split("."):
        m = re.match(r"^([^\[]+)(\[(\d+)\])?$", parte)
        if not m:
            return None
        cur = (cur or {}).get(m.group(1)) if isinstance(cur, dict) else None
        if m.group(3) is not None:
            if isinstance(cur, list) and len(cur) > int(m.group(3)):
                cur = cur[int(m.group(3))]
            else:
                return None
    return cur


def render_pdf_overlay(contenido: dict[str, Any], titulo: str, fecha: datetime) -> bytes:
    """Reproduce la plantilla de 6 páginas (PNG de fondo) con los datos sobrepuestos."""
    from weasyprint import HTML

    om = _overlay_map()
    w_mm, h_mm = om.get("page_size_mm", [179, 228])
    debug = settings.export_debug_grid

    paginas = []
    for i in range(1, 7):
        uri = _png_uri(i)
        bg = f'<img class="bg" src="{uri}"/>' if uri else ""
        campos = []
        for ruta, pos in (om["pages"].get(str(i), {}) or {}).items():
            val = _resolver(contenido, ruta)
            if val in (None, "", "-.-", []):
                continue
            estilo = (f"top:{pos['top']}%;left:{pos['left']}%;"
                      f"font-size:{pos.get('size', 9)}px;"
                      + (f"width:{pos['width']}%;" if pos.get("width") else "")
                      + ("font-weight:bold;" if pos.get("bold") else ""))
            campos.append(f'<div class="f" style="{estilo}">{val}</div>')
        grid = '<div class="grid"></div>' if debug else ""
        paginas.append(f'<div class="pg">{bg}{grid}{"".join(campos)}</div>')

    grid_css = (
        ".grid{position:absolute;inset:0;background-image:"
        "repeating-linear-gradient(to right,rgba(255,0,0,.25) 0 1px,transparent 1px 10%),"
        "repeating-linear-gradient(to bottom,rgba(255,0,0,.25) 0 1px,transparent 1px 10%);}"
        if debug else ""
    )
    html = f"""<!doctype html><html><head><meta charset="utf-8"><style>
      @page {{ size: {w_mm}mm {h_mm}mm; margin: 0; }}
      * {{ margin:0; padding:0; box-sizing:border-box; }}
      .pg {{ position:relative; width:{w_mm}mm; height:{h_mm}mm; page-break-after:always; overflow:hidden; }}
      .pg:last-child {{ page-break-after:auto; }}
      .bg {{ position:absolute; inset:0; width:100%; height:100%; }}
      .f {{ position:absolute; font-family:"DejaVu Sans",Arial,sans-serif; color:#111; line-height:1.05; }}
      {grid_css}
    </style></head><body>{"".join(paginas)}</body></html>"""
    return HTML(string=html).write_pdf()


def render_pdf_flow(contenido: dict[str, Any], titulo: str, fecha: datetime) -> bytes:
    from weasyprint import HTML

    ctx = {
        "titulo": titulo,
        "fecha": fecha.strftime("%d%b%Y").upper(),
        "anio": fecha.year,
        "resumen_ejecutivo": contenido.get("resumen_ejecutivo", []),
        "situacion": contenido.get("situacion", {}),
        "asuntos_criticos": contenido.get("asuntos_criticos", []),
        "proyeccion_24_72h": contenido.get("proyeccion_24_72h", ""),
        "personal": contenido.get("personal", {}),
        "inteligencia": contenido.get("inteligencia", {}),
        "operaciones": contenido.get("operaciones", {}),
        "logistica": contenido.get("logistica", {}),
        "donut_personal": _donut_personal(contenido.get("personal", {})),
        "bar_macrozonas": _bar_macrozonas(contenido.get("personal", {})),
        "bar_traslados": _bar_traslados(contenido.get("operaciones", {})),
    }
    html = _env.get_template("briefing.html.j2").render(**ctx)
    return HTML(string=html).write_pdf()


def render_pdf(contenido: dict[str, Any], titulo: str, fecha: datetime) -> bytes:
    """Despacha al modo configurado: overlay (pixel-perfect) o flow (reconstrucción)."""
    if settings.export_style.lower() == "flow":
        return render_pdf_flow(contenido, titulo, fecha)
    return render_pdf_overlay(contenido, titulo, fecha)


def render_texto(contenido: dict[str, Any], titulo: str, fecha: datetime) -> bytes:
    L = [f"{titulo}  ({fecha.strftime('%d-%m-%Y')})", "=" * 60, "", "RESUMEN EJECUTIVO:"]
    for b in contenido.get("resumen_ejecutivo", []) or ["-.-"]:
        L.append(f"  • {b}")
    sit = contenido.get("situacion") or {}
    L += ["", "SITUACIÓN:", "  " + (sit.get("resumen") or "-.-")]
    for a in sit.get("aspectos", []) or []:
        L.append(f"  • {a}")
    L += ["", "ASUNTOS CRÍTICOS:"]
    for a in contenido.get("asuntos_criticos", []) or []:
        L.append(f"  - {a.get('asunto','-.-')} | {a.get('impacto','-.-')} | {a.get('responsable','-.-')}")
    L += ["", "PROYECCIÓN 24-72H:", "  " + (contenido.get("proyeccion_24_72h") or "-.-")]
    return "\n".join(L).encode("utf-8")


def render_word(contenido: dict[str, Any], titulo: str, fecha: datetime) -> bytes:
    import docx

    doc = docx.Document()
    doc.add_heading(titulo, level=0)
    doc.add_paragraph(fecha.strftime("%d-%m-%Y"))

    doc.add_heading("Resumen ejecutivo", level=1)
    for b in contenido.get("resumen_ejecutivo", []) or ["-.-"]:
        doc.add_paragraph(b, style="List Bullet")

    sit = contenido.get("situacion") or {}
    doc.add_heading("Situación", level=1)
    doc.add_paragraph(sit.get("resumen") or "-.-")
    for a in sit.get("aspectos", []) or []:
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Asuntos críticos", level=1)
    ac = contenido.get("asuntos_criticos", []) or []
    if ac:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Asunto", "Impacto", "Responsable"
        for a in ac:
            c = t.add_row().cells
            c[0].text, c[1].text, c[2].text = a.get("asunto", "-.-"), a.get("impacto", "-.-"), a.get("responsable", "-.-")

    doc.add_heading("Proyección 24-72h", level=1)
    doc.add_paragraph(contenido.get("proyeccion_24_72h") or "-.-")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def render(contenido: dict[str, Any], formato: str, titulo: str, fecha: datetime) -> tuple[bytes, str, str]:
    """Devuelve (bytes, content_type, extension) según el formato."""
    formato = formato.upper()
    if formato == "PDF":
        return render_pdf(contenido, titulo, fecha), "application/pdf", "pdf"
    if formato == "WORD":
        return (
            render_word(contenido, titulo, fecha),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "docx",
        )
    return render_texto(contenido, titulo, fecha), "text/plain; charset=utf-8", "txt"
