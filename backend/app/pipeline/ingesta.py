"""
Filtro 1 — INGESTA (RF-001): de binario heterogéneo a texto plano.

Soporta PDF, Word, Excel, correos (.msg/.eml) y bitácoras/texto. El binario
llega desde MinIO (ya descifrado por storage.get_bytes).
"""
from __future__ import annotations

import io
import os
import subprocess
import tempfile

# Firmas de archivo: .docx es un ZIP (PK\x03\x04); .doc 97-2003 es OLE2.
_ZIP_MAGIC = b"PK\x03\x04"
_OLE2_MAGIC = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"


def _texto_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return "\n".join((p.extract_text() or "") for p in reader.pages)


def _texto_doc_binario(data: bytes) -> str:
    """Extrae texto de un .doc binario (Word 97-2003) con antiword (RF-001)."""
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as tmp:
        tmp.write(data)
        ruta = tmp.name
    try:
        out = subprocess.run(
            ["antiword", ruta],
            capture_output=True, timeout=60,
        )
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout.decode("utf-8", errors="ignore")
        # Fallback final: volcado de texto legible del binario.
        return data.decode("latin-1", errors="ignore")
    except (FileNotFoundError, subprocess.SubprocessError):
        return data.decode("latin-1", errors="ignore")
    finally:
        os.unlink(ruta)


def _texto_word(data: bytes, nombre: str = "") -> str:
    # .doc binario (OLE2) o extensión .doc -> antiword; python-docx solo lee .docx (ZIP).
    es_ole2 = data[:8] == _OLE2_MAGIC
    es_zip = data[:4] == _ZIP_MAGIC
    if es_ole2 or (nombre.lower().endswith(".doc") and not es_zip):
        return _texto_doc_binario(data)

    import docx

    doc = docx.Document(io.BytesIO(data))
    partes = [p.text for p in doc.paragraphs]
    for tabla in doc.tables:
        for fila in tabla.rows:
            partes.append("\t".join(c.text for c in fila.cells))
    return "\n".join(partes)


def _texto_excel(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    lineas: list[str] = []
    for hoja in wb.worksheets:
        lineas.append(f"# Hoja: {hoja.title}")
        for fila in hoja.iter_rows(values_only=True):
            celdas = [str(c) for c in fila if c is not None]
            if celdas:
                lineas.append("\t".join(celdas))
    return "\n".join(lineas)


def _texto_correo(data: bytes, nombre: str) -> str:
    ext = os.path.splitext(nombre)[1].lower()
    if ext == ".msg":
        import extract_msg

        with tempfile.NamedTemporaryFile(suffix=".msg", delete=False) as tmp:
            tmp.write(data)
            ruta = tmp.name
        try:
            msg = extract_msg.openMsg(ruta)
            return f"De: {msg.sender}\nAsunto: {msg.subject}\nFecha: {msg.date}\n\n{msg.body}"
        finally:
            os.unlink(ruta)
    # .eml u otros
    try:
        import mailparser

        mail = mailparser.parse_from_bytes(data)
        return f"De: {mail.from_}\nAsunto: {mail.subject}\nFecha: {mail.date}\n\n{mail.body}"
    except Exception:
        return data.decode("utf-8", errors="ignore")


def _texto_plano(data: bytes) -> str:
    return data.decode("utf-8", errors="ignore")


def extraer_texto(data: bytes, tipo: str, nombre_archivo: str) -> str:
    """Devuelve el texto plano de un documento según su tipo (RF-001)."""
    tipo = (tipo or "OTRO").upper()
    if tipo == "PDF":
        return _texto_pdf(data)
    if tipo == "WORD":
        return _texto_word(data, nombre_archivo)
    if tipo == "EXCEL":
        return _texto_excel(data)
    if tipo == "CORREO":
        return _texto_correo(data, nombre_archivo)
    return _texto_plano(data)
