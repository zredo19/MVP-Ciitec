#!/usr/bin/env python3
"""
Generador de corpus sintético en español para demo y RNF-004 (50 docs < 3 min).

Crea documentos operacionales heterogéneos (PDF, Word, Excel, correo, bitácora)
con hechos plausibles (unidades, fechas, lugares, incidentes, personal, meteo,
logística) para alimentar la extracción y la síntesis del briefing institucional.

Salida: ./corpus/  (50 documentos + 5 "curados" más ricos).

Dependencias: python-docx, openpyxl y (opcional) fpdf2 para PDF.
    pip install python-docx openpyxl fpdf2
"""
from __future__ import annotations

import os
import random
from datetime import datetime, timedelta

random.seed(42)

OUT = os.path.join(os.path.dirname(__file__), "..", "corpus")
os.makedirs(OUT, exist_ok=True)

UNIDADES = ["I División de Ejército", "JAF Arica y Parinacota", "JAF Tarapacá", "JAF Antofagasta",
            "JDN Biobío", "JDN Araucanía", "Fuerza de Tarea Andes", "BAE O'Higgins"]
LUGARES = ["Arica", "Iquique", "Antofagasta", "Calama", "Concepción", "Temuco", "Punta Arenas", "Putre"]
INCIDENTES = [
    ("Incendio forestal", "Apoyo a combate de incendio forestal con medios de ingenieros"),
    ("Deslizamiento", "Despeje de ruta tras deslizamiento de tierra por lluvias"),
    ("Sismo", "Evaluación de daños y apoyo a la población tras sismo"),
    ("Aluvión", "Evacuación preventiva y rescate por aluvión en quebrada"),
    ("Marejadas", "Refuerzo de borde costero ante marejadas anormales"),
]
ESTADOS = ["ABIERTO", "EN_CURSO", "CERRADO"]
METEO = [
    "Se pronostican entre 10 y 15 mm de precipitaciones y vientos de 20 a 30 km/h en la zona central.",
    "Cielos despejados con temperaturas bajo cero en el altiplano; riesgo de heladas.",
    "Sistema frontal aproximándose por el sur, oleaje en aumento en el litoral.",
]


def _fecha() -> datetime:
    return datetime(2026, 5, 1) + timedelta(days=random.randint(0, 20), hours=random.randint(0, 23))


def _parrafo_hecho() -> str:
    u = random.choice(UNIDADES)
    l = random.choice(LUGARES)
    tipo, desc = random.choice(INCIDENTES)
    f = _fecha()
    est = random.choice(ESTADOS)
    fuerza = random.randint(20, 180)
    return (f"El {f.strftime('%d-%m-%Y')} la unidad {u} ejecutó en {l} la siguiente actividad: "
            f"{desc} ({tipo}). Estado: {est}. Fuerza empeñada: {fuerza} efectivos. "
            f"Responsable: Comandante de {u}.")


def _texto_doc(n_hechos: int) -> str:
    L = ["PARTE OPERACIONAL DIARIO — USO RESERVADO", ""]
    for _ in range(n_hechos):
        L.append("- " + _parrafo_hecho())
    L += ["", "SITUACIÓN METEOROLÓGICA:", random.choice(METEO),
          "", f"PERSONAL: OF {random.randint(40,80)}, SOF {random.randint(100,200)}, "
          f"ECP {random.randint(200,400)}, TOTAL {random.randint(400,800)}."]
    return "\n".join(L)


def gen_txt(path: str, n: int):
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(_texto_doc(n))


def gen_eml(path: str, n: int):
    f = _fecha()
    cuerpo = _texto_doc(n)
    eml = (f"From: operaciones@ejercito.cl\nTo: comando@ejercito.cl\n"
           f"Subject: Parte operacional {f.strftime('%d%b%Y')}\nDate: {f.strftime('%a, %d %b %Y %H:%M:%S -0400')}\n\n{cuerpo}")
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(eml)


def gen_docx(path: str, n: int):
    import docx
    doc = docx.Document()
    doc.add_heading("Parte Operacional — Reservado", level=1)
    for linea in _texto_doc(n).splitlines():
        doc.add_paragraph(linea)
    doc.save(path)


def gen_xlsx(path: str, n: int):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Operacionalidad"
    ws.append(["Unidad", "Total", "NOP", "Campaña", "Combate", "Fecha"])
    for _ in range(n):
        ws.append([random.choice(UNIDADES), random.randint(100, 400), random.randint(10, 90),
                   random.randint(10, 50), random.randint(10, 50), _fecha().strftime("%d-%m-%Y")])
    wb.save(path)


def gen_pdf(path: str, n: int) -> bool:
    try:
        from fpdf import FPDF
    except ImportError:
        return False
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", size=11)
    for linea in _texto_doc(n).splitlines():
        pdf.multi_cell(0, 6, linea.encode("latin-1", "replace").decode("latin-1"))
    pdf.output(path)
    return True


def main():
    plan = (["pdf"] * 15) + (["docx"] * 15) + (["xlsx"] * 10) + (["eml"] * 5) + (["txt"] * 5)
    random.shuffle(plan)
    creados, sin_pdf = 0, 0
    for i, fmt in enumerate(plan, 1):
        base = os.path.join(OUT, f"doc_{i:02d}")
        n = random.randint(3, 7)
        if fmt == "pdf":
            if not gen_pdf(base + ".pdf", n):
                gen_txt(base + ".txt", n); sin_pdf += 1
        elif fmt == "docx":
            gen_docx(base + ".docx", n)
        elif fmt == "xlsx":
            gen_xlsx(base + ".xlsx", n)
        elif fmt == "eml":
            gen_eml(base + ".eml", n)
        else:
            gen_txt(base + ".txt", n)
        creados += 1

    # 5 "curados" más ricos (10-14 hechos)
    for j in range(1, 6):
        gen_docx(os.path.join(OUT, f"curado_{j}.docx"), random.randint(10, 14))
        creados += 1

    print(f"Corpus generado en {os.path.abspath(OUT)}: {creados} documentos.")
    if sin_pdf:
        print(f"  (fpdf2 no instalado: {sin_pdf} PDFs se generaron como .txt; `pip install fpdf2` para PDFs reales)")


if __name__ == "__main__":
    main()
